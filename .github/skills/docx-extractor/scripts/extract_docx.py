#!/usr/bin/env python3
"""
Word文档(DOCX)文本提取工具
提取Word文档中的段落文本和表格内容，保存为纯文本格式
"""

import argparse
from pathlib import Path

from docx import Document
from rich.console import Console
from rich.table import Table as RichTable

console = Console()


def extract_docx_content(docx_path: Path, output_path: Path | None = None) -> str:
    """
    提取DOCX文档的文本和表格内容
    
    Args:
        docx_path: DOCX文件路径
        output_path: 输出文件路径（可选）
    
    Returns:
        提取的文本内容
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        console.print(f"[red]❌ 无法打开文档: {e}[/red]")
        raise
    
    console.print(f"[green]✓ 文档打开成功[/green]")
    
    content_lines = []
    content_lines.append(f"文档: {docx_path.name}")
    content_lines.append("=" * 70)
    content_lines.append("")
    
    # 统计信息
    paragraph_count = 0
    table_count = 0
    
    # 遍历文档的所有元素（段落和表格）
    for element in doc.element.body:
        # 处理段落
        if element.tag.endswith('p'):
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if para and para.text.strip():
                content_lines.append(para.text)
                paragraph_count += 1
        
        # 处理表格
        elif element.tag.endswith('tbl'):
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break
            
            if table:
                table_count += 1
                content_lines.append("")
                content_lines.append(f"【表格 {table_count}】")
                content_lines.append("-" * 70)
                
                # 提取表格内容
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ')
                        row_data.append(cell_text)
                    content_lines.append(f"行 {row_idx + 1}: {' | '.join(row_data)}")
                
                content_lines.append("-" * 70)
                content_lines.append("")
    
    # 合并所有内容
    full_content = "\n".join(content_lines)
    
    # 显示统计信息
    stats_table = RichTable(title="提取统计")
    stats_table.add_column("指标", style="cyan")
    stats_table.add_column("值", style="green")
    stats_table.add_row("段落数", str(paragraph_count))
    stats_table.add_row("表格数", str(table_count))
    stats_table.add_row("总字符数", str(len(full_content)))
    console.print(stats_table)
    
    # 保存到文件
    if output_path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_content)
        console.print(f"[green]✓ 提取完成！[/green]")
        console.print(f"[green]✓ 结果已保存到: {output_path}[/green]")
    
    return full_content


def main():
    parser = argparse.ArgumentParser(
        description="提取Word文档(DOCX)中的文本和表格内容",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 提取DOCX文档内容并保存
  python extract_docx.py document.docx -o output.txt
  
  # 使用uv运行
  uv run python extract_docx.py document.docx -o output.txt
        """
    )
    
    parser.add_argument(
        'docx_file',
        type=Path,
        help='要提取的DOCX文件路径'
    )
    
    parser.add_argument(
        '-o', '--output',
        type=Path,
        default=None,
        help='输出文本文件路径（默认: 与输入文件同名的.txt文件）'
    )
    
    args = parser.parse_args()
    
    # 检查输入文件是否存在
    if not args.docx_file.exists():
        console.print(f"[red]❌ 文件不存在: {args.docx_file}[/red]")
        return 1
    
    # 确定输出路径
    if args.output is None:
        args.output = args.docx_file.with_suffix('.txt')
    
    try:
        extract_docx_content(args.docx_file, args.output)
        return 0
    except Exception as e:
        console.print(f"[red]❌ 提取失败: {e}[/red]")
        return 1


if __name__ == '__main__':
    exit(main())
